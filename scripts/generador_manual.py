#!/usr/bin/env python
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from playwright.sync_api import Page, TimeoutError as PlaywrightTimeoutError, sync_playwright


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_APP_URL = "https://ancloraadvisorai-ten.vercel.app/"
DEFAULT_TEMPLATE = ROOT / "docs" / "Free_Simple_ SAAS_Agreement_Template.docx"
DEFAULT_PLAN = ROOT / "manual.screenshots.yml"
DEFAULT_OUTPUT = ROOT / "manual-assets" / "MANUAL_USUARIO_CODEX.docx"
DEFAULT_SCREENSHOT_DIR = ROOT / "manual-assets" / "screenshots"
DEFAULT_COVER = ROOT / "manual-assets" / "cover" / "manual-cover.png"
DEFAULT_REPORT = ROOT / "report.json"
AUTH_STATE = ROOT / ".auth" / "state.json"


DEFAULT_SCREENSHOT_PLAN: dict[str, Any] = {
    "meta": {
        "app_name": "Anclora Advisor AI",
        "app_url": DEFAULT_APP_URL,
        "run_mode": "remote",
        "viewport": {"width": 1440, "height": 900},
        "root_selector_default": "main",
    },
    "screenshots": [
        {
            "id": "login",
            "title": "Pantalla de acceso",
            "path": "/login",
            "filename": "01-login.png",
            "root_selector": "main",
            "auth_required": False,
        },
        {
            "id": "dashboard_chat",
            "title": "Dashboard principal y chat RAG",
            "path": "/dashboard/chat",
            "filename": "02-dashboard-chat.png",
            "root_selector": "main",
        },
        {
            "id": "dashboard_fiscal",
            "title": "Módulo fiscal",
            "path": "/dashboard/fiscal",
            "filename": "03-dashboard-fiscal.png",
            "root_selector": "main",
        },
        {
            "id": "dashboard_laboral",
            "title": "Módulo laboral",
            "path": "/dashboard/laboral",
            "filename": "04-dashboard-laboral.png",
            "root_selector": "main",
        },
        {
            "id": "dashboard_facturacion",
            "title": "Módulo de facturación",
            "path": "/dashboard/facturacion",
            "filename": "05-dashboard-facturacion.png",
            "root_selector": "main",
        },
        {
            "id": "dashboard_alertas",
            "title": "Centro de alertas",
            "path": "/dashboard/alertas",
            "filename": "06-dashboard-alertas.png",
            "root_selector": "main",
        },
        {
            "id": "alerts_modal",
            "title": "Panel rápido de notificaciones",
            "path": "/dashboard/chat",
            "filename": "07-alerts-modal.png",
            "root_selector": "body",
            "actions": [
                {"type": "click_role", "role": "button", "name": "Notificaciones"},
                {"type": "wait_for_text", "text": "Notificaciones"},
            ],
        },
    ],
}


REPO_SOURCES = [
    "src/components/auth/LoginForm.tsx",
    "src/components/layout/DashboardNav.tsx",
    "src/components/layout/DashboardTopbar.tsx",
    "src/components/layout/GeneralAlertCenter.tsx",
    "docs/FISCAL_MODULE_OPERATIVO.md",
    "docs/LABOR_MODULE_OPERATIVO.md",
    "docs/INVOICE_MODULE_OPERATIVO.md",
    "docs/GENERAL_ALERTS.md",
]


@dataclass
class ScreenshotResult:
    id: str
    title: str
    path: str
    filename: str
    absolute_path: str | None = None
    status: str = "pending"
    notes: str | None = None


@dataclass
class BuilderReport:
    app_url: str
    run_mode: str
    generated_files: list[str] = field(default_factory=list)
    screenshot_index: list[dict[str, Any]] = field(default_factory=list)
    sections_added: list[str] = field(default_factory=list)
    gaps: list[dict[str, str]] = field(default_factory=list)
    sources_used: list[str] = field(default_factory=list)
    glossary: list[dict[str, str]] = field(default_factory=list)
    assumptions: list[str] = field(default_factory=list)

    def add_gap(self, category: str, detail: str) -> None:
        self.gaps.append({"category": category, "detail": detail})


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Genera un manual de usuario DOCX a partir del repo y la app remota.")
    parser.add_argument("--app-url", default=DEFAULT_APP_URL)
    parser.add_argument("--run-mode", default="remote", choices=["remote"])
    parser.add_argument("--template", default=str(DEFAULT_TEMPLATE))
    parser.add_argument("--plan", default=str(DEFAULT_PLAN))
    parser.add_argument("--output-docx", default=str(DEFAULT_OUTPUT))
    parser.add_argument("--screenshots-dir", default=str(DEFAULT_SCREENSHOT_DIR))
    parser.add_argument("--report", default=str(DEFAULT_REPORT))
    parser.add_argument("--plan-only", action="store_true", help="No abre Playwright; genera plan y reporte.")
    parser.add_argument("--allow-partial", action="store_true", help="Genera manual y reporte aunque falten capturas.")
    return parser.parse_args()


def ensure_directory(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_plan(plan_path: Path, report: BuilderReport) -> dict[str, Any]:
    if not plan_path.exists():
        plan_path.write_text(yaml.safe_dump(DEFAULT_SCREENSHOT_PLAN, sort_keys=False, allow_unicode=True), encoding="utf-8")
        report.assumptions.append("manual.screenshots.yml no existía; se generó un plan base con rutas principales del dashboard.")
        return DEFAULT_SCREENSHOT_PLAN
    with plan_path.open("r", encoding="utf-8") as handle:
        return yaml.safe_load(handle) or DEFAULT_SCREENSHOT_PLAN


def save_plan(plan_path: Path, plan: dict[str, Any]) -> None:
    plan_path.write_text(yaml.safe_dump(plan, sort_keys=False, allow_unicode=True), encoding="utf-8")


def infer_root_selector(page: Page, requested: str | None) -> str:
    candidates = [requested, "main", '[role="main"]', "body"]
    seen: set[str] = set()
    for candidate in candidates:
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        try:
            locator = page.locator(candidate).first
            locator.wait_for(state="visible", timeout=4_000)
            return candidate
        except PlaywrightTimeoutError:
            continue
    return "body"


def hide_transient_ui(page: Page) -> None:
    page.add_style_tag(
        content="""
        [data-sonner-toast], [data-radix-toast-viewport], .toast, [role="status"], [aria-live="polite"] {
          visibility: hidden !important;
          opacity: 0 !important;
        }
        """
    )


def wait_stable(page: Page, root_selector: str) -> None:
    page.wait_for_load_state("networkidle")
    page.locator(root_selector).first.wait_for(state="visible", timeout=12_000)


def is_logged_in(page: Page) -> bool:
    return "/login" not in page.url and page.locator("text=Asesoría RAG").count() > 0


def login_if_needed(page: Page, app_url: str, report: BuilderReport) -> None:
    page.goto(f"{app_url.rstrip('/')}/dashboard/chat", wait_until="domcontentloaded")
    try:
        wait_stable(page, "body")
    except PlaywrightTimeoutError:
        pass

    if is_logged_in(page):
        return

    user = os.getenv("MANUAL_TEST_USER")
    password = os.getenv("MANUAL_TEST_PASS")
    if not user or not password:
        raise RuntimeError("Las variables MANUAL_TEST_USER y MANUAL_TEST_PASS son obligatorias para autenticar en remoto.")

    page.goto(f"{app_url.rstrip('/')}/login", wait_until="domcontentloaded")
    wait_stable(page, "main")
    page.locator("#email").fill(user)
    page.locator("#password").fill(password)
    page.get_by_role("button", name="Entrar al dashboard").click()
    try:
        page.wait_for_url(re.compile(r".*/dashboard/.*"), timeout=20_000)
    except PlaywrightTimeoutError as exc:
        page.wait_for_timeout(2_000)
        if "/login" in page.url:
            body_text = page.locator("body").inner_text()
            auth_error = None
            for pattern in (
                "Invalid login credentials",
                "Credenciales inválidas",
                "Invalid Refresh Token",
            ):
                if pattern in body_text:
                    auth_error = pattern
                    break
            if auth_error:
                raise RuntimeError(f"Autenticación remota rechazada: {auth_error}") from exc
        raise RuntimeError("No se completó la redirección al dashboard tras el login.") from exc
    wait_stable(page, "main")
    report.assumptions.append("Se realizó autenticación UI y se guardó storageState en ./.auth/state.json.")


def run_actions(page: Page, actions: list[dict[str, Any]]) -> None:
    for action in actions:
        action_type = action.get("type")
        if action_type == "click_role":
            page.get_by_role(action["role"], name=action["name"]).click()
        elif action_type == "wait_for_text":
            page.get_by_text(action["text"], exact=action.get("exact", True)).first.wait_for(state="visible", timeout=10_000)
        elif action_type == "click_selector":
            page.locator(action["selector"]).first.click()
        elif action_type == "wait_for_selector":
            page.locator(action["selector"]).first.wait_for(state="visible", timeout=10_000)


def capture_screenshots(
    plan: dict[str, Any],
    app_url: str,
    screenshots_dir: Path,
    report: BuilderReport,
) -> list[ScreenshotResult]:
    ensure_directory(screenshots_dir)
    ensure_directory(AUTH_STATE.parent)
    viewport = plan.get("meta", {}).get("viewport", {"width": 1440, "height": 900})
    default_root_selector = plan.get("meta", {}).get("root_selector_default", "main")
    results: list[ScreenshotResult] = []

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        storage_state = str(AUTH_STATE) if AUTH_STATE.exists() else None
        context = browser.new_context(
            viewport=viewport,
            storage_state=storage_state,
            locale="es-ES",
            color_scheme="dark",
        )
        page = context.new_page()
        hide_transient_ui(page)

        try:
            login_if_needed(page, app_url, report)
            context.storage_state(path=str(AUTH_STATE))
        except Exception as exc:  # noqa: BLE001
            browser.close()
            raise RuntimeError(str(exc)) from exc

        for step in plan.get("screenshots", []):
            shot = ScreenshotResult(
                id=step["id"],
                title=step["title"],
                path=step["path"],
                filename=step["filename"],
            )
            try:
                target_url = f"{app_url.rstrip('/')}{step['path']}"
                page.goto(target_url, wait_until="domcontentloaded")
                root_selector = infer_root_selector(page, step.get("root_selector", default_root_selector))
                if root_selector != step.get("root_selector"):
                    step["root_selector"] = root_selector
                wait_stable(page, root_selector)
                if step.get("actions"):
                    run_actions(page, step["actions"])
                    wait_stable(page, root_selector if root_selector != "body" else "body")
                screenshot_path = screenshots_dir / step["filename"]
                page.screenshot(path=str(screenshot_path), full_page=False)
                shot.absolute_path = str(screenshot_path)
                shot.status = "captured"
            except Exception as exc:  # noqa: BLE001
                shot.status = "failed"
                shot.notes = str(exc)
                report.add_gap("SCREENSHOT_CAPTURE", f"{step['id']}: {exc}")
            results.append(shot)

        browser.close()
    return results


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def ensure_caption_style(doc: Document) -> str:
    try:
        doc.styles["Caption"]
        return "Caption"
    except KeyError:
        style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.size = Pt(9)
        style.font.italic = True
        return "Caption"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Paragraph")


def add_screenshot(doc: Document, caption_style: str, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(6.5))
    doc.add_paragraph(caption, style=caption_style)


def find_font(preferred: list[str]) -> str | None:
    font_dirs = [
        Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts",
        Path("/usr/share/fonts"),
        Path("/usr/local/share/fonts"),
    ]
    for font_dir in font_dirs:
        for name in preferred:
            candidate = font_dir / name
            if candidate.exists():
                return str(candidate)
    return None


def build_manual_cover(cover_path: Path) -> Path:
    ensure_directory(cover_path.parent)
    width, height = 1800, 2550
    image = Image.new("RGB", (width, height), "#0b1830")
    draw = ImageDraw.Draw(image)

    for y in range(height):
        mix = y / max(height - 1, 1)
        r = int(11 + (21 - 11) * mix)
        g = int(24 + (47 - 24) * mix)
        b = int(48 + (78 - 48) * mix)
        draw.line((0, y, width, y), fill=(r, g, b))

    glow = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    glow_draw = ImageDraw.Draw(glow)
    glow_draw.ellipse((180, 220, 1620, 1540), fill=(38, 197, 173, 40))
    glow_draw.ellipse((350, 1180, 1450, 2300), fill=(90, 170, 255, 26))
    glow = glow.filter(ImageFilter.GaussianBlur(120))
    image = Image.alpha_composite(image.convert("RGBA"), glow)

    overlay = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    overlay_draw = ImageDraw.Draw(overlay)
    panel_bounds = (130, 250, width - 130, height - 250)
    overlay_draw.rounded_rectangle(
        panel_bounds,
        radius=70,
        fill=(8, 22, 44, 170),
        outline=(70, 110, 150, 110),
        width=3,
    )
    overlay_draw.line((220, 1420, width - 220, 1420), fill=(54, 191, 174, 90), width=2)
    image = Image.alpha_composite(image, overlay)

    title_font_path = find_font(["Georgia.ttf", "georgia.ttf", "GARA.TTF", "timesbd.ttf"])
    body_font_path = find_font(["Arial.ttf", "arial.ttf", "segoeui.ttf"])
    title_font = ImageFont.truetype(title_font_path, 104) if title_font_path else ImageFont.load_default()
    subtitle_font = ImageFont.truetype(body_font_path, 46) if body_font_path else ImageFont.load_default()
    tag_font = ImageFont.truetype(body_font_path, 52) if body_font_path else ImageFont.load_default()

    logo_path = ROOT / "public" / "brand" / "logo-Advisor_1.png"
    if logo_path.exists():
        logo = Image.open(logo_path).convert("RGBA")
        logo.thumbnail((330, 330))
        logo_x = (width - logo.width) // 2
        image.alpha_composite(logo, (logo_x, 430))

    text_layer = ImageDraw.Draw(image)
    title = "Manual de Usuario\nde Anclora Advisor AI"
    title_box = text_layer.multiline_textbbox((0, 0), title, font=title_font, spacing=18, align="center")
    title_w = title_box[2] - title_box[0]
    title_x = (width - title_w) / 2
    text_layer.multiline_text(
        (title_x, 860),
        title,
        font=title_font,
        fill=(236, 242, 250, 255),
        spacing=18,
        align="center",
    )

    subtitle = "Guía operativa, funcional y visual"
    subtitle_box = text_layer.textbbox((0, 0), subtitle, font=subtitle_font)
    subtitle_x = (width - (subtitle_box[2] - subtitle_box[0])) / 2
    text_layer.text((subtitle_x, 1180), subtitle, font=subtitle_font, fill=(156, 203, 233, 255))

    tag = "By Anclora Group"
    tag_box = text_layer.textbbox((0, 0), tag, font=tag_font)
    tag_x = (width - (tag_box[2] - tag_box[0])) / 2
    text_layer.text((tag_x, 2010), tag, font=tag_font, fill=(41, 203, 176, 255))

    image.convert("RGB").save(cover_path, format="PNG", optimize=True)
    return cover_path


def add_cover_page(doc: Document, cover_path: Path, caption_style: str) -> None:
    if not cover_path.exists():
        return
    doc.add_picture(str(cover_path), width=Inches(6.5))
    doc.add_paragraph("Portada. Manual de Usuario de Anclora Advisor AI.", style=caption_style)
    doc.add_page_break()


def build_glossary() -> list[dict[str, str]]:
    return [
        {"term": "RAG", "definition": "Consulta asistida sobre la base documental fiscal, laboral y de mercado."},
        {"term": "Serie", "definition": "Prefijo operativo usado para numerar facturas."},
        {"term": "Rectificativa", "definition": "Factura que corrige una factura origen y queda vinculada a ella."},
        {"term": "Recordatorio recurrente", "definition": "Plantilla periódica que genera alertas futuras automáticamente."},
        {"term": "SLA", "definition": "Fecha objetivo de seguimiento en una mitigación laboral."},
        {"term": "Verifactu", "definition": "Integración prevista para envío de facturas al sistema Verifactu."},
    ]


def assemble_manual(
    template_path: Path,
    output_docx: Path,
    screenshot_results: list[ScreenshotResult],
    report: BuilderReport,
) -> None:
    doc = Document(str(template_path))
    clear_document(doc)
    caption_style = ensure_caption_style(doc)
    cover_path = build_manual_cover(DEFAULT_COVER)

    captured = {result.id: result for result in screenshot_results if result.absolute_path}
    glossary = build_glossary()
    report.glossary = glossary
    report.generated_files.append(str(cover_path.relative_to(ROOT)))

    add_cover_page(doc, cover_path, caption_style)

    add_heading(doc, "Manual de usuario - Anclora Advisor AI", 1)
    doc.add_paragraph(
        "Documento generado a partir del repositorio y de la validación visual de la aplicación remota. "
        "Las funciones no verificadas en UI quedan marcadas como 'Pendiente de validación'.",
        style="Normal",
    )
    doc.add_paragraph(
        f"Fecha de generación: {datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M %Z')}",
        style="Normal",
    )

    sections = [
        "Getting Started",
        "Navigation map",
        "Task-based guides",
        "Account/Profile",
        "Glossary",
        "FAQ",
        "Troubleshooting",
        "Support",
    ]
    report.sections_added = sections

    add_heading(doc, "Getting Started", 1)
    doc.add_paragraph("1. Accede a la URL de la plataforma y autentícate con tu correo y contraseña.", style="Normal")
    doc.add_paragraph("2. Tras iniciar sesión entrarás en el dashboard y podrás navegar por los módulos laterales.", style="Normal")
    doc.add_paragraph("3. Desde la barra superior puedes cambiar idioma, tema visual y revisar alertas.", style="Normal")
    if "login" in captured:
        add_screenshot(doc, caption_style, Path(captured["login"].absolute_path), "Figura 1. Pantalla de acceso.")
    if "dashboard_chat" in captured:
        add_screenshot(doc, caption_style, Path(captured["dashboard_chat"].absolute_path), "Figura 2. Dashboard principal y chat RAG.")

    add_heading(doc, "Navigation map", 1)
    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    header = table.rows[0].cells
    header[0].text = "Ruta"
    header[1].text = "Módulo"
    header[2].text = "Uso principal"
    rows = [
        ("/dashboard/chat", "Chat", "Consulta normativa con historial persistido."),
        ("/dashboard/fiscal", "Fiscal", "Calendario, alertas y plantillas fiscales."),
        ("/dashboard/laboral", "Laboral", "Evaluaciones de riesgo y mitigaciones."),
        ("/dashboard/facturacion", "Facturación", "Facturas, cobros, PDF, Verifactu e importación."),
        ("/dashboard/alertas", "Alertas", "Centro transversal de avisos y recurrencias."),
        ("/dashboard/admin", "Admin", "Solo para roles admin. Pendiente de validación de permisos por perfil."),
    ]
    for route, module, purpose in rows:
        cells = table.add_row().cells
        cells[0].text = route
        cells[1].text = module
        cells[2].text = purpose

    add_heading(doc, "Task-based guides", 1)
    add_heading(doc, "1. Consultar normativa en el chat", 2)
    add_bullets(
        doc,
        [
            "Abre Chat desde el menú lateral.",
            "Escribe una consulta y pulsa 'Consultar'.",
            "Usa el historial persistido para reabrir conversaciones anteriores.",
        ],
    )
    add_heading(doc, "2. Gestionar obligaciones fiscales", 2)
    add_bullets(
        doc,
        [
            "Crea alertas fiscales y define estado de tramitación.",
            "Usa plantillas mensuales, trimestrales o anuales para obligaciones recurrentes.",
            "Revisa el resumen de pendientes, vencidas y modelos fiscales activos.",
        ],
    )
    add_heading(doc, "3. Gestionar riesgo laboral", 2)
    add_bullets(
        doc,
        [
            "Registra una evaluación de riesgo desde el módulo Laboral.",
            "Asigna responsable, SLA, checklist y evidencias a cada mitigación.",
            "Filtra por escenario, responsable y estado para priorizar seguimiento.",
        ],
    )
    add_heading(doc, "4. Emitir y controlar facturas", 2)
    add_bullets(
        doc,
        [
            "Crea o edita facturas desde un formulario único.",
            "Revisa serie, numeración, estado, cobros parciales y exportaciones.",
            "Usa la vista PDF, el envío por correo y las acciones de Verifactu cuando estén configuradas.",
            "La importación desde PDF/imagen está soportada, pero el OCR/VLM queda pendiente de validación remota.",
        ],
    )
    add_heading(doc, "5. Gestionar avisos y recordatorios", 2)
    add_bullets(
        doc,
        [
            "Abre la campana del header para revisar notificaciones rápidas.",
            "Usa filtros por categoría y paginación por bloques de cuatro alertas.",
            "Accede al Centro de Alertas para ver recordatorios recurrentes y seguimiento completo.",
        ],
    )
    for key, caption in [
        ("dashboard_fiscal", "Figura 3. Módulo fiscal."),
        ("dashboard_laboral", "Figura 4. Módulo laboral."),
        ("dashboard_facturacion", "Figura 5. Módulo de facturación."),
        ("dashboard_alertas", "Figura 6. Centro de alertas."),
        ("alerts_modal", "Figura 7. Modal rápido de notificaciones."),
    ]:
        if key in captured:
            add_screenshot(doc, caption_style, Path(captured[key].absolute_path), caption)

    add_heading(doc, "Account/Profile", 1)
    add_bullets(
        doc,
        [
            "Inicio de sesión, creación de cuenta y recuperación de contraseña están visibles en la pantalla de acceso.",
            "Cambio de idioma y tema visual están disponibles en la barra superior del dashboard.",
            "Edición de perfil, avatar o preferencias avanzadas: Pendiente de validación.",
            "Cierre de sesión disponible desde la parte inferior del menú lateral.",
        ],
    )

    add_heading(doc, "Glossary", 1)
    for item in glossary:
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.add_run(f"{item['term']}: ").bold = True
        paragraph.add_run(item["definition"])

    add_heading(doc, "FAQ", 1)
    faq_pairs = [
        ("¿Qué ocurre si no veo el dashboard tras iniciar sesión?", "Revisa que las cookies y la sesión de Supabase sigan vigentes. Si vuelve a aparecer /login, inicia sesión de nuevo."),
        ("¿Puedo usar la plataforma en varios módulos sin salir del dashboard?", "Sí. La navegación lateral permite pasar entre chat, fiscal, laboral, facturación y alertas."),
        ("¿Cómo recibo avisos del navegador?", "Desde la campana del header puedes conceder permiso al navegador y activar las notificaciones."),
        ("¿Dónde edito mis datos personales?", "Pendiente de validación: no se ha verificado una pantalla específica de perfil editable."),
    ]
    for question, answer in faq_pairs:
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.add_run(question + " ").bold = True
        paragraph.add_run(answer)

    add_heading(doc, "Troubleshooting", 1)
    add_bullets(
        doc,
        [
            "Si la sesión expira, vuelve a /login y autentícate de nuevo.",
            "Si una captura o una acción falla por cambios en la UI, revisa manual.screenshots.yml y corrige el selector o el texto accesible.",
            "Si el envío SMTP o Verifactu no está configurado, la operación puede quedar en cola o fallar con error operativo.",
            "Si una función no aparece para tu usuario, puede depender del rol o de configuración pendiente.",
        ],
    )

    add_heading(doc, "Support", 1)
    doc.add_paragraph(
        "Canal de soporte visible en la aplicación: Pendiente de validación. "
        "Mientras no exista un canal explicitado en UI, documenta internamente quién gestiona incidencias funcionales, acceso y datos.",
        style="Normal",
    )

    ensure_directory(output_docx.parent)
    doc.save(str(output_docx))
    report.generated_files.append(str(output_docx.relative_to(ROOT)))


def write_report(report_path: Path, report: BuilderReport) -> None:
    payload = {
        "app_url": report.app_url,
        "run_mode": report.run_mode,
        "generated_files": report.generated_files,
        "screenshot_index": report.screenshot_index,
        "sections_added": report.sections_added,
        "gaps": report.gaps,
        "sources_used": report.sources_used,
        "glossary": report.glossary,
        "assumptions": report.assumptions,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }
    report_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def main() -> int:
    args = parse_args()
    plan_path = Path(args.plan)
    output_docx = Path(args.output_docx)
    screenshots_dir = Path(args.screenshots_dir)
    report_path = Path(args.report)
    template_path = Path(args.template)

    report = BuilderReport(app_url=args.app_url, run_mode=args.run_mode)
    report.sources_used = REPO_SOURCES.copy()

    if not template_path.exists():
        report.add_gap("TEMPLATE_MISSING", f"No se encontró la plantilla DOCX en {template_path}")
        write_report(report_path, report)
        print(f"Template not found: {template_path}", file=sys.stderr)
        return 1

    plan = load_plan(plan_path, report)
    screenshot_results: list[ScreenshotResult] = []

    if not args.plan_only:
        try:
            screenshot_results = capture_screenshots(plan, args.app_url, screenshots_dir, report)
            save_plan(plan_path, plan)
        except Exception as exc:  # noqa: BLE001
            report.add_gap("AUTH_OR_PLAYWRIGHT", str(exc))
            if not args.allow_partial:
                write_report(report_path, report)
                print(str(exc), file=sys.stderr)
                return 1
    else:
        report.assumptions.append("Ejecución en modo plan-only: no se capturaron pantallas reales.")

    report.screenshot_index = [
        {
            "id": result.id,
            "title": result.title,
            "path": result.path,
            "filename": result.filename,
            "status": result.status,
            "file": str(Path(result.absolute_path).relative_to(ROOT)) if result.absolute_path else None,
            "notes": result.notes,
        }
        for result in screenshot_results
    ]

    if screenshot_results or args.allow_partial or args.plan_only:
        assemble_manual(template_path, output_docx, screenshot_results, report)
    else:
        report.add_gap("MANUAL_SKIPPED", "No se generó DOCX porque falló la autenticación/captura y no se solicitó allow-partial.")

    report.generated_files.extend(
        [
            str(plan_path.relative_to(ROOT)),
            str(report_path.relative_to(ROOT)),
        ]
    )
    write_report(report_path, report)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
