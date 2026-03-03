#!/usr/bin/env python
"""
manual_builder.py — Anclora Advisor AI User Manual Generator
=============================================================

Generates a complete end-user DOCX manual from:
  - A Word template (respecting its styles)
  - Real screenshots taken from the live remote app via Playwright
  - A curated screenshot plan (manual.screenshots.yml)

Usage
-----
  python scripts/manual_builder.py [OPTIONS]

Options
  --template PATH     Word template (.docx). Default: docs/Free_Simple_ SAAS_Agreement_Template.docx
  --plan PATH         Screenshot plan YAML. Default: manual.screenshots.yml
  --out PATH          Output DOCX path. Default: artifacts/AncoraAdvisorAI_Manual.docx
  --report PATH       Report JSON path. Default: report.json
  --assets-dir PATH   Screenshot output dir. Default: manual-assets/screenshots
  --auth-state PATH   Playwright storageState. Default: .auth/state.json
  --headless          Run browser headless (default: headed for debugging)
  --dry-run           Take screenshots only; skip DOCX assembly
  --screenshots-only  Same as --dry-run
  --skip-screenshots  Use existing screenshots; only assemble DOCX

Environment variables (REQUIRED for auth — never hardcode)
  MANUAL_TEST_USER   App login e-mail
  MANUAL_TEST_PASS   App login password

Security contract
  - Credentials are read exclusively from env vars.
  - They are never logged, stored in code, YAML, or any report file.
  - The auth state file (.auth/state.json) is gitignored.
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from playwright.sync_api import sync_playwright, TimeoutError as PWTimeoutError

# ─── Logging ──────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("manual_builder")

# ─── Constants ────────────────────────────────────────────────────────────────

REPO_ROOT = Path(__file__).parent.parent
DEFAULT_TEMPLATE = REPO_ROOT / "docs" / "Free_Simple_ SAAS_Agreement_Template.docx"
DEFAULT_PLAN = REPO_ROOT / "manual.screenshots.yml"
DEFAULT_OUT = REPO_ROOT / "artifacts" / "AncoraAdvisorAI_Manual.docx"
DEFAULT_REPORT = REPO_ROOT / "report.json"
DEFAULT_ASSETS = REPO_ROOT / "manual-assets" / "screenshots"
DEFAULT_AUTH_STATE = REPO_ROOT / ".auth" / "state.json"

APP_URL = "https://ancloraadvisorai-ten.vercel.app"
LOGIN_URL = f"{APP_URL}/login"
POST_LOGIN_PATH = "/dashboard/chat"

NETWORK_IDLE_TIMEOUT = 20_000   # ms
NAV_TIMEOUT = 60_000             # ms

# ─── .env.local loader ────────────────────────────────────────────────────────

def _load_env_local() -> None:
    """
    Load variables from .env.local into os.environ (only if not already set).
    Does NOT override existing shell environment variables.
    Supports KEY=VALUE and KEY="VALUE" (strips surrounding quotes).
    """
    env_file = REPO_ROOT / ".env.local"
    if not env_file.exists():
        return
    for line in env_file.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, raw_val = line.partition("=")
        key = key.strip()
        val = raw_val.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = val
    log.debug("Loaded .env.local")


# ─── Manual content ───────────────────────────────────────────────────────────

MANUAL_SECTIONS: list[dict[str, Any]] = [
    {
        "id": "intro",
        "heading": "Introducción",
        "level": 1,
        "content": (
            "Anclora Advisor AI es una plataforma de asesoría inteligente diseñada para "
            "profesionales autónomos españoles. Integra inteligencia artificial con bases "
            "de conocimiento especializadas en fiscalidad, riesgos laborales y mercado "
            "inmobiliario en Baleares, ofreciendo orientación personalizada en tiempo real."
        ),
    },
    {
        "id": "getting_started",
        "heading": "Primeros pasos",
        "level": 1,
        "content": None,
        "subsections": [
            {
                "heading": "Requisitos previos",
                "content": (
                    "• Navegador web moderno (Chrome, Firefox, Safari, Edge).\n"
                    "• Conexión a internet.\n"
                    "• Cuenta de usuario proporcionada por el administrador."
                ),
            },
            {
                "heading": "Acceso e inicio de sesión",
                "content": (
                    "1. Abra su navegador y acceda a https://ancloraadvisorai-ten.vercel.app\n"
                    "2. Introduzca su correo electrónico y contraseña.\n"
                    "3. Pulse «Iniciar sesión». Si las credenciales son correctas, será "
                    "redirigido al módulo Chat.\n\n"
                    "Si no recuerda su contraseña, contacte con el administrador de la plataforma."
                ),
                "screenshot_id": "01-login",   # 01-login.png
            },
        ],
    },
    {
        "id": "navigation",
        "heading": "Mapa de navegación",
        "level": 1,
        "content": (
            "La plataforma se organiza en cinco módulos principales accesibles desde la "
            "barra lateral izquierda:\n\n"
            "• 💬 Chat — Asistente de asesoría con IA y búsqueda RAG.\n"
            "• 🧾 Fiscal — Gestión de alertas y plazos tributarios.\n"
            "• ⚖️ Laboral — Evaluación de riesgos y acciones de mitigación.\n"
            "• 🧮 Facturación — Emisión y seguimiento de facturas.\n"
            "• 🔔 Alertas — Centro de alertas generales y recordatorios.\n\n"
            "La barra lateral puede contraerse pulsando el botón de flecha para ganar "
            "espacio de trabajo. En modo contraído se muestran únicamente los iconos."
        ),
        "subsections": [
            {
                "heading": "Barra lateral expandida",
                "content": "La barra lateral muestra el nombre de cada módulo, su subtítulo descriptivo y un indicador de estado activo.",
                "screenshot_id": "02-dashboard-chat",
            },
            {
                "heading": "Barra lateral contraída",
                "content": "En modo compacto, los módulos se identifican únicamente por su icono. Pase el cursor sobre cada icono para ver su nombre.",
                "screenshot_id": "02-dashboard-chat",
            },
        ],
    },
    {
        "id": "chat",
        "heading": "Chat Asesor (IA)",
        "level": 1,
        "content": (
            "El módulo Chat es el punto de entrada principal a la inteligencia artificial "
            "de Anclora. Permite formular preguntas en lenguaje natural sobre fiscalidad, "
            "situaciones laborales o mercado inmobiliario. Las respuestas se basan en "
            "fuentes verificadas mediante búsqueda RAG (Retrieval-Augmented Generation)."
        ),
        "subsections": [
            {
                "heading": "Iniciar una conversación",
                "content": (
                    "1. Navegue a Chat desde la barra lateral.\n"
                    "2. Escriba su consulta en el campo de texto inferior.\n"
                    "3. Pulse Enter o el botón de envío.\n"
                    "4. El asistente responderá con orientación y, cuando proceda, "
                    "incluirá acciones sugeridas (botones de acción rápida)."
                ),
                "screenshot_id": "02-dashboard-chat",
            },
            {
                "heading": "Historial de conversaciones",
                "content": (
                    "El panel lateral izquierdo del Chat muestra todas sus conversaciones "
                    "anteriores. Puede:\n"
                    "• Pulsar sobre cualquier conversación para continuarla.\n"
                    "• Crear una nueva conversación con el botón «+».\n"
                    "• Renombrar conversaciones para facilitar su organización."
                ),
                "screenshot_id": "07-alerts-modal",
            },
            {
                "heading": "Acciones sugeridas",
                "content": (
                    "Tras algunas respuestas, el asistente propone acciones concretas "
                    "(p. ej. «Crear alerta fiscal», «Ver facturas pendientes»). "
                    "Al pulsarlas, el sistema le redirige al módulo correspondiente o "
                    "ejecuta la acción directamente."
                ),
            },
            {
                "heading": "Citaciones y fuentes",
                "content": (
                    "Las respuestas que se basan en normativa o jurisprudencia incluyen "
                    "citas al pie. Revise siempre las fuentes antes de tomar decisiones "
                    "con implicaciones legales o fiscales."
                ),
            },
        ],
    },
    {
        "id": "fiscal",
        "heading": "Módulo Fiscal",
        "level": 1,
        "content": (
            "El módulo Fiscal centraliza la gestión de alertas tributarias, plazos de "
            "declaración y flujos de trabajo fiscales para el autónomo. Cubre IVA, IRPF, "
            "RETA, declaraciones informativas y otras obligaciones de la AEAT y la ATIB."
        ),
        "subsections": [
            {
                "heading": "Listado de alertas fiscales",
                "content": (
                    "La pantalla principal muestra todas sus alertas ordenadas por fecha "
                    "límite. Puede filtrar por estado (pendiente, en curso, completada) "
                    "y por tipo de impuesto."
                ),
                "screenshot_id": "03-dashboard-fiscal",
            },
            {
                "heading": "Crear una nueva alerta fiscal",
                "content": (
                    "1. Pulse «Nueva alerta».\n"
                    "2. Seleccione el tipo de alerta (IVA, IRPF, RETA, etc.).\n"
                    "3. Indique la fecha límite y la prioridad.\n"
                    "4. (Opcional) Asigne régimen fiscal y modelo tributario.\n"
                    "5. Pulse «Guardar»."
                ),
                "screenshot_id": "03-dashboard-fiscal",
            },
            {
                "heading": "Plantillas fiscales",
                "content": (
                    "Las plantillas permiten crear alertas recurrentes automáticamente "
                    "(p. ej. Modelo 303 trimestral). Configure la recurrencia y el día "
                    "de vencimiento para que el sistema genere las alertas de forma "
                    "automática cada período."
                ),
            },
            {
                "heading": "Registro de auditoría",
                "content": (
                    "Cada cambio en una alerta fiscal queda registrado en el timeline de "
                    "auditoría. Esto permite rastrear quién modificó qué y cuándo, "
                    "fundamental para documentar la diligencia debida ante la AEAT."
                ),
            },
        ],
    },
    {
        "id": "laboral",
        "heading": "Módulo Laboral",
        "level": 1,
        "content": (
            "El módulo Laboral ayuda al autónomo a evaluar y gestionar los riesgos "
            "derivados de su situación laboral: pluriactividad, compatibilidad con contrato "
            "por cuenta ajena, riesgo de despido disciplinario, y otros escenarios de "
            "transición."
        ),
        "subsections": [
            {
                "heading": "Evaluaciones de riesgo",
                "content": (
                    "Cada evaluación describe un escenario laboral y le asigna una "
                    "puntuación de riesgo (0–100) y nivel (bajo, medio, alto, crítico). "
                    "El asistente genera recomendaciones específicas para mitigar cada riesgo."
                ),
                "screenshot_id": "04-dashboard-laboral",
            },
            {
                "heading": "Crear una evaluación de riesgo",
                "content": (
                    "1. Pulse «Nueva evaluación».\n"
                    "2. Describa el escenario laboral.\n"
                    "3. Indique la puntuación de riesgo estimada o déjela en 0 para "
                    "que el sistema la calcule.\n"
                    "4. Pulse «Guardar»."
                ),
                "screenshot_id": "04-dashboard-laboral",
            },
            {
                "heading": "Acciones de mitigación",
                "content": (
                    "Para cada evaluación puede crear acciones de mitigación: tareas "
                    "concretas con fecha límite, responsable y lista de verificación. "
                    "El estado de cada acción (pendiente, en curso, cerrada) se refleja "
                    "en el panel lateral."
                ),
            },
            {
                "heading": "Evidencias y cierre",
                "content": (
                    "Al cerrar una acción de mitigación, puede adjuntar notas de cierre "
                    "y enlaces a evidencias documentales (correos, resoluciones, contratos). "
                    "Estas evidencias quedan almacenadas en el historial de auditoría."
                ),
            },
        ],
    },
    {
        "id": "facturacion",
        "heading": "Módulo de Facturación",
        "level": 1,
        "content": (
            "El módulo de Facturación permite emitir, gestionar y hacer seguimiento de "
            "facturas. Incluye soporte para IVA, retención de IRPF, rectificativas, "
            "pagos parciales y compatibilidad con VeriFactu (Ley Antifraude 2022)."
        ),
        "subsections": [
            {
                "heading": "Listado de facturas",
                "content": (
                    "La pantalla principal muestra todas sus facturas con filtros por "
                    "estado (borrador, emitida, cobrada, anulada), tipo (ordinaria, "
                    "rectificativa) y rango de fechas."
                ),
                "screenshot_id": "05-dashboard-facturacion",
            },
            {
                "heading": "Emitir una nueva factura",
                "content": (
                    "1. Pulse «Nueva factura».\n"
                    "2. Complete los datos del cliente (nombre y NIF).\n"
                    "3. Introduzca el importe base, tipo de IVA y retención de IRPF.\n"
                    "4. Indique la fecha de emisión y la serie.\n"
                    "5. (Opcional) Añada el correo del receptor para envío automático.\n"
                    "6. Pulse «Guardar»."
                ),
                "screenshot_id": "05-dashboard-facturacion",
            },
            {
                "heading": "Registro de pagos",
                "content": (
                    "Una vez emitida la factura, puede registrar pagos totales o "
                    "parciales indicando el método de pago, la referencia y la fecha. "
                    "El sistema actualiza automáticamente el estado de la factura."
                ),
            },
            {
                "heading": "VeriFactu",
                "content": (
                    "Las facturas emitidas se marcan con su estado VeriFactu "
                    "(pendiente, enviada, verificada). Esta funcionalidad cumple con los "
                    "requisitos de la Ley Antifraude 11/2021 y el Real Decreto VeriFactu 2024. "
                    "Pendiente de validación: integración directa con AEAT."
                ),
            },
        ],
    },
    {
        "id": "alertas",
        "heading": "Centro de Alertas",
        "level": 1,
        "content": (
            "El Centro de Alertas es el punto central para gestionar notificaciones "
            "generales que no encajan en los módulos Fiscal o Laboral. Permite crear "
            "alertas con categoría, prioridad, fecha límite y recordatorios recurrentes."
        ),
        "subsections": [
            {
                "heading": "Vista general de alertas",
                "content": (
                    "Las alertas se organizan por categoría (fiscal, laboral, facturación) "
                    "y se pueden filtrar. Cada alerta muestra su prioridad (crítica, alta, "
                    "media, baja) y la fecha límite."
                ),
                "screenshot_id": "06-dashboard-alertas",
            },
            {
                "heading": "Crear una alerta con recurrencia",
                "content": (
                    "1. Pulse «Nueva alerta».\n"
                    "2. Rellene título, mensaje, categoría, prioridad y fecha límite.\n"
                    "3. Active la recurrencia (diaria, semanal, mensual, anual) si procede.\n"
                    "4. Indique los días de antelación para el recordatorio.\n"
                    "5. Pulse «Guardar»."
                ),
                "screenshot_id": "07-alerts-modal",
            },
            {
                "heading": "Recordatorios recurrentes",
                "content": (
                    "Los recordatorios recurrentes se generan automáticamente en base a "
                    "la recurrencia configurada. Aparecen en el panel derecho del Centro "
                    "de Alertas y pueden editarse o eliminarse individualmente."
                ),
            },
        ],
    },
    {
        "id": "account",
        "heading": "Cuenta y preferencias",
        "level": 1,
        "content": None,
        "subsections": [
            {
                "heading": "Barra superior",
                "content": (
                    "La barra superior de la aplicación muestra su correo electrónico, "
                    "el rol activo y los controles de preferencias."
                ),
                "screenshot_id": "02-dashboard-chat",
            },
            {
                "heading": "Idioma",
                "content": (
                    "La plataforma soporta español (ES) e inglés (EN). Cambie el idioma "
                    "desde el selector en la barra superior. El cambio es inmediato y "
                    "persiste entre sesiones."
                ),
            },
            {
                "heading": "Tema visual",
                "content": (
                    "Dispone de modo claro y modo oscuro. Cambie el tema desde el icono "
                    "de luna/sol en la barra superior. El tema seleccionado se guarda "
                    "automáticamente."
                ),
            },
            {
                "heading": "Cerrar sesión",
                "content": (
                    "Para cerrar sesión, pulse «Cerrar sesión» en la parte inferior de "
                    "la barra lateral. Su sesión quedará invalidada de forma segura."
                ),
            },
        ],
    },
    {
        "id": "glossary",
        "heading": "Glosario",
        "level": 1,
        "content": None,
        "glossary": [
            ("AEAT", "Agencia Estatal de Administración Tributaria. Organismo que gestiona los impuestos en España."),
            ("ATIB", "Agència Tributària de les Illes Balears. Organismo tributario autonómico de Baleares."),
            ("IAE", "Impuesto de Actividades Económicas. Obligatorio para empresas con facturación > 1 M€; exento para autónomos."),
            ("IRPF", "Impuesto sobre la Renta de las Personas Físicas. Los autónomos declaran rendimientos mediante Modelo 130 o 131."),
            ("IVA", "Impuesto sobre el Valor Añadido. Se liquida trimestralmente mediante el Modelo 303."),
            ("Pluriactividad", "Situación en la que un trabajador está dado de alta simultáneamente en el RETA y en el Régimen General de la Seguridad Social."),
            ("RAG", "Retrieval-Augmented Generation. Técnica de IA que combina búsqueda en bases de datos vectoriales con generación de texto."),
            ("RETA", "Régimen Especial de Trabajadores Autónomos. Sistema de Seguridad Social para trabajadores por cuenta propia."),
            ("VeriFactu", "Sistema de facturación electrónica verificable exigido por la Ley Antifraude 11/2021 a partir de julio 2025."),
        ],
    },
    {
        "id": "faq",
        "heading": "Preguntas frecuentes (FAQ)",
        "level": 1,
        "content": None,
        "faq": [
            (
                "¿Las respuestas del Chat tienen valor legal?",
                "No. Las respuestas son orientativas y se basan en fuentes documentales verificadas. "
                "Para decisiones con implicaciones legales o fiscales, consulte siempre con un "
                "asesor profesional colegiado.",
            ),
            (
                "¿Puedo importar facturas desde otro programa?",
                "Actualmente la importación masiva no está disponible. Pendiente de validación.",
            ),
            (
                "¿Cómo sé si una alerta fiscal se ha completado?",
                "Abra la alerta y cambie su estado a «Completada». El sistema registrará la "
                "fecha y hora del cambio en el timeline de auditoría.",
            ),
            (
                "¿El asistente tiene en cuenta mi situación específica?",
                "El Chat usa su historial de conversaciones para contextualizar las respuestas, "
                "pero no tiene acceso a su contabilidad externa ni a sus datos de la AEAT.",
            ),
            (
                "¿Cómo actualizo mi contraseña?",
                "Contacte con el administrador de la plataforma. La gestión de credenciales se "
                "realiza desde el panel de administración (rol admin). Pendiente de validación.",
            ),
            (
                "¿La plataforma funciona en móvil?",
                "La interfaz es responsive, pero está optimizada para escritorio (≥ 1024 px de ancho).",
            ),
        ],
    },
    {
        "id": "troubleshooting",
        "heading": "Solución de problemas",
        "level": 1,
        "content": None,
        "subsections": [
            {
                "heading": "No puedo iniciar sesión",
                "content": (
                    "• Verifique que el correo y la contraseña son correctos.\n"
                    "• Compruebe que no tiene el teclado en otro idioma (p. ej. mayúsculas activadas).\n"
                    "• Borre las cookies del navegador e inténtelo de nuevo.\n"
                    "• Contacte con el administrador si el problema persiste."
                ),
            },
            {
                "heading": "El Chat no responde",
                "content": (
                    "• Compruebe su conexión a internet.\n"
                    "• Espere unos segundos: la primera respuesta puede tardar hasta 10 s.\n"
                    "• Recargue la página (F5) y reintente.\n"
                    "• Si el error persiste, abra una nueva conversación."
                ),
            },
            {
                "heading": "Una alerta no aparece en el listado",
                "content": (
                    "• Verifique el filtro de estado activo: es posible que la alerta esté "
                    "filtrada como «Completada».\n"
                    "• Limpie todos los filtros con el botón correspondiente."
                ),
            },
            {
                "heading": "La factura no se ha enviado por correo",
                "content": (
                    "• Confirme que el campo «Correo del receptor» estaba relleno antes de guardar.\n"
                    "• El envío por correo puede tardar hasta 5 minutos.\n"
                    "• Revise la carpeta de spam del destinatario."
                ),
            },
        ],
    },
    {
        "id": "support",
        "heading": "Soporte y contacto",
        "level": 1,
        "content": (
            "Para cualquier incidencia técnica o consulta sobre la plataforma, contacte "
            "con el equipo de soporte de Anclora:\n\n"
            "• Correo: soporte@anclora.ai  (Pendiente de validación)\n"
            "• Horario: lunes a viernes, 9:00–18:00 (hora peninsular)\n\n"
            "Incluya en su mensaje: descripción del problema, módulo afectado, "
            "nombre de usuario y capturas de pantalla si procede."
        ),
    },
]

GLOSSARY = {
    entry["id"]: entry.get("glossary", [])
    for entry in MANUAL_SECTIONS
    if "glossary" in entry
}

# ─── Argument parsing ─────────────────────────────────────────────────────────

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Generate Anclora Advisor AI user manual (DOCX) with real screenshots.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    p.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE, help="Word template (.docx)")
    p.add_argument("--plan", type=Path, default=DEFAULT_PLAN, help="Screenshot plan YAML")
    p.add_argument("--out", type=Path, default=DEFAULT_OUT, help="Output DOCX path")
    p.add_argument("--report", type=Path, default=DEFAULT_REPORT, help="Report JSON path")
    p.add_argument("--assets-dir", type=Path, default=DEFAULT_ASSETS, help="Screenshot output directory")
    p.add_argument("--auth-state", type=Path, default=DEFAULT_AUTH_STATE, help="Playwright storageState file")
    p.add_argument("--headless", action="store_true", help="Run browser headless")
    p.add_argument("--dry-run", "--screenshots-only", action="store_true", help="Take screenshots only; skip DOCX assembly")
    p.add_argument("--skip-screenshots", action="store_true", help="Skip screenshots; use existing ones")
    return p.parse_args()

# ─── YAML plan loading ────────────────────────────────────────────────────────

def _normalize_plan(raw: dict) -> dict:
    """
    Normalize both YAML schemas into a single internal format.

    Schema A (v1, original):
      app_url / viewport / auth / screenshots[].{id, url, caption, wait_selector, pre_actions, full_page}

    Schema B (user-curated):
      meta.{app_url, viewport, root_selector_default} /
      screenshots[].{id, path, filename, title, root_selector, actions[].{type, role, name, text}}
    """
    meta = raw.get("meta", {})
    app_url = (meta.get("app_url") or raw.get("app_url") or APP_URL).rstrip("/")
    vp = meta.get("viewport") or raw.get("viewport") or {"width": 1440, "height": 900}
    default_sel = meta.get("root_selector_default") or "main"

    normalized_shots = []
    for shot in raw.get("screenshots", []):
        filename = shot.get("filename", "")
        stem = Path(filename).stem if filename else shot.get("id", "unknown")
        url = shot.get("path") or shot.get("url") or "/"
        caption = shot.get("title") or shot.get("caption") or stem
        wait_sel = shot.get("root_selector") or shot.get("wait_selector") or default_sel
        full_page = shot.get("full_page", False)
        section = shot.get("section", "")
        auth_required = shot.get("auth_required", True)

        # Normalize actions (Schema B) → pre_actions (Schema A internal)
        pre_actions: list[dict] = []
        for act in shot.get("actions", shot.get("pre_actions", [])):
            act_type = act.get("type") or act.get("action", "")
            if act_type == "click_role":
                pre_actions.append({
                    "action": "click_role",
                    "role": act.get("role", "button"),
                    "name": act.get("name", ""),
                    "timeout": act.get("timeout", 8000),
                    "optional": act.get("optional", True),
                })
            elif act_type == "wait_for_text":
                pre_actions.append({
                    "action": "wait_for_text",
                    "text": act.get("text", ""),
                    "timeout": act.get("timeout", 8000),
                    "optional": act.get("optional", True),
                })
            elif act_type in ("click", "fill", "wait", "goto"):
                pre_actions.append({**act, "action": act_type})
            else:
                pre_actions.append(act)

        normalized_shots.append({
            "id": stem,
            "url": url,
            "caption": caption,
            "wait_selector": wait_sel,
            "pre_actions": pre_actions,
            "full_page": full_page,
            "section": section,
            "auth_required": auth_required,
        })

    auth_cfg = raw.get("auth", {})
    return {
        "app_url": app_url,
        "viewport": vp,
        "auth": auth_cfg,
        "screenshots": normalized_shots,
    }


def load_plan(plan_path: Path) -> dict[str, Any]:
    if not plan_path.exists():
        log.warning("Screenshot plan not found at %s — generating minimal default", plan_path)
        default_plan = _default_plan()
        plan_path.write_text(
            yaml.dump(default_plan, allow_unicode=True, sort_keys=False),
            encoding="utf-8",
        )
        log.info("Default plan written to %s", plan_path)
        return default_plan
    with plan_path.open(encoding="utf-8") as fh:
        raw = yaml.safe_load(fh)
    return _normalize_plan(raw)


def _default_plan() -> dict[str, Any]:
    return {
        "version": "1",
        "app_url": APP_URL,
        "viewport": {"width": 1440, "height": 900},
        "auth": {
            "state_file": ".auth/state.json",
            "login_url": "/login",
            "post_login_url": "/dashboard/chat",
        },
        "screenshots": [
            {"id": "01-login", "section": "getting_started", "url": "/login",
             "caption": "Pantalla de inicio de sesión", "wait_selector": "input[type='email']"},
            {"id": "02-chat-empty", "section": "chat", "url": "/dashboard/chat",
             "caption": "Vista principal del Chat", "wait_selector": "main"},
            {"id": "04-fiscal-overview", "section": "fiscal", "url": "/dashboard/fiscal",
             "caption": "Módulo Fiscal", "wait_selector": "main"},
            {"id": "06-laboral-overview", "section": "laboral", "url": "/dashboard/laboral",
             "caption": "Módulo Laboral", "wait_selector": "main"},
            {"id": "08-facturacion-overview", "section": "facturacion", "url": "/dashboard/facturacion",
             "caption": "Módulo Facturación", "wait_selector": "main"},
            {"id": "10-alertas-overview", "section": "alertas", "url": "/dashboard/alertas",
             "caption": "Centro de Alertas", "wait_selector": "main"},
        ],
    }

# ─── Auth helpers ─────────────────────────────────────────────────────────────

def _read_credentials() -> tuple[str, str]:
    user = os.environ.get("MANUAL_TEST_USER", "")
    pwd = os.environ.get("MANUAL_TEST_PASS", "")
    if not user or not pwd:
        log.error(
            "Missing credentials. Set MANUAL_TEST_USER and MANUAL_TEST_PASS "
            "environment variables before running."
        )
        sys.exit(1)
    return user, pwd


def ensure_auth(page: Any, plan: dict, auth_state: Path, assets_dir: Path) -> bool:
    """
    Ensures the browser is authenticated.
    Re-uses storageState if valid; re-logs if missing or expired.
    Returns True if auth succeeded.
    """
    login_path = plan.get("auth", {}).get("login_url", "/login")
    post_login = plan.get("auth", {}).get("post_login_url", "/dashboard/chat")
    base_url = plan.get("app_url", APP_URL)

    # Check if we're already authenticated by navigating to the post-login page.
    try:
        page.goto(base_url + post_login, timeout=NAV_TIMEOUT, wait_until="domcontentloaded")
        page.wait_for_timeout(2000)
        if post_login.split("/")[-1] in page.url or "dashboard" in page.url:
            log.info("Session valid — reusing stored auth state.")
            return True
    except Exception:
        pass

    log.info("Session invalid or missing — performing UI login.")
    user, pwd = _read_credentials()

    page.goto(base_url + login_path, timeout=NAV_TIMEOUT, wait_until="domcontentloaded")

    # Fill email
    email_input = page.locator("input[type='email'], input[name='email'], input[placeholder*='mail' i]").first
    email_input.wait_for(timeout=15_000)
    email_input.fill(user)

    # Fill password
    pwd_input = page.locator("input[type='password']").first
    pwd_input.wait_for(timeout=10_000)
    pwd_input.fill(pwd)

    # Submit
    submit = page.locator("button[type='submit'], button:has-text('Iniciar'), button:has-text('Login'), button:has-text('Entrar')").first
    submit.click()

    # Wait for redirect to dashboard
    try:
        page.wait_for_url(f"**{post_login}**", timeout=30_000)
    except PWTimeoutError:
        # Fallback: wait for main
        page.locator("main").wait_for(timeout=20_000)

    log.info("Login successful.")
    # Save state
    auth_state.parent.mkdir(parents=True, exist_ok=True)
    page.context.storage_state(path=str(auth_state))
    log.info("Auth state saved to %s", auth_state)
    return True

# ─── Screenshot engine ────────────────────────────────────────────────────────

def _hide_toasts(page: Any) -> None:
    """Try to hide toast/notification overlays before capture."""
    try:
        page.evaluate(
            """() => {
                const selectors = [
                    '[data-sonner-toast]', '[data-radix-toast-viewport]',
                    '.toast', '.Toastify', '[role="alert"]',
                    '[data-state="open"][data-type]',
                ];
                selectors.forEach(sel => {
                    document.querySelectorAll(sel).forEach(el => {
                        el.style.display = 'none';
                    });
                });
            }"""
        )
    except Exception:
        pass


def _execute_pre_actions(page: Any, plan_shot: dict, base_url: str) -> list[str]:
    """Execute pre_actions for a screenshot entry. Returns list of warnings."""
    warnings: list[str] = []
    for act in plan_shot.get("pre_actions", []):
        action = act.get("action")
        selector = act.get("selector", "")
        timeout = act.get("timeout", 8000)
        optional = act.get("optional", False)
        try:
            if action == "click":
                el = page.locator(selector).first
                el.wait_for(timeout=timeout, state="visible")
                el.click()
                page.wait_for_timeout(1500)
            elif action == "click_role":
                role = act.get("role", "button")
                name = act.get("name", "")
                el = page.get_by_role(role, name=name).first
                el.wait_for(timeout=timeout, state="visible")
                el.click()
                page.wait_for_timeout(1500)
            elif action == "wait_for_text":
                text = act.get("text", "")
                page.get_by_text(text).first.wait_for(timeout=timeout, state="visible")
            elif action == "fill":
                el = page.locator(selector).first
                el.wait_for(timeout=timeout)
                el.fill(act.get("value", ""))
            elif action == "wait":
                page.wait_for_timeout(timeout)
            elif action == "goto":
                page.goto(base_url + act.get("url", ""), timeout=NAV_TIMEOUT)
        except Exception as exc:
            msg = f"Pre-action '{action}' on '{selector or act.get('name', '')}' failed: {exc}"
            if optional:
                log.warning("  [optional] %s", msg)
                warnings.append(msg)
            else:
                raise RuntimeError(msg) from exc
    return warnings


def take_screenshots(
    plan: dict,
    auth_state: Path,
    assets_dir: Path,
    headless: bool,
) -> tuple[list[dict], list[str]]:
    """
    Navigates the live app and captures all screenshots defined in the plan.
    Returns (index, gaps) where index is a list of screenshot metadata dicts
    and gaps is a list of warning strings.
    """
    assets_dir.mkdir(parents=True, exist_ok=True)
    base_url: str = plan.get("app_url", APP_URL)
    vp = plan.get("viewport", {"width": 1440, "height": 900})
    screenshots_plan: list[dict] = plan.get("screenshots", [])

    index: list[dict] = []
    gaps: list[str] = []

    storage_state_arg = str(auth_state) if auth_state.exists() else None

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=headless, channel="chrome")
        ctx_kwargs: dict[str, Any] = {
            "viewport": vp,
            "locale": "es-ES",
            "timezone_id": "Europe/Madrid",
            "color_scheme": "light",
        }
        if storage_state_arg:
            ctx_kwargs["storage_state"] = storage_state_arg

        ctx = browser.new_context(**ctx_kwargs)
        page = ctx.new_page()

        # Authenticate
        ensure_auth(page, plan, auth_state, assets_dir)

        # Capture each screenshot
        for shot in screenshots_plan:
            sid = shot.get("id", "unknown")
            url = shot.get("url", "/")
            caption = shot.get("caption", sid)
            wait_sel = shot.get("wait_selector", "main")
            full_page = shot.get("full_page", False)
            out_path = assets_dir / f"{sid}.png"

            log.info("Screenshot [%s]: %s", sid, base_url + url)
            try:
                page.goto(base_url + url, timeout=NAV_TIMEOUT, wait_until="domcontentloaded")

                # Detect auth redirect — session may have expired or been cleared
                # (e.g. visiting /login can invalidate Supabase tokens in some setups)
                if shot.get("auth_required", True) and "/login" in page.url:
                    log.warning("  Auth redirect detected for %s — re-authenticating", sid)
                    ensure_auth(page, plan, auth_state, assets_dir)
                    page.goto(base_url + url, timeout=NAV_TIMEOUT, wait_until="domcontentloaded")

                # Wait for network idle
                try:
                    page.wait_for_load_state("networkidle", timeout=NETWORK_IDLE_TIMEOUT)
                except PWTimeoutError:
                    log.warning("  networkidle timeout for %s — continuing anyway", sid)

                # Wait for stable selector
                try:
                    page.locator(wait_sel).first.wait_for(state="visible", timeout=15_000)
                except PWTimeoutError:
                    w = f"wait_selector '{wait_sel}' not found for {sid}"
                    log.warning("  %s", w)
                    gaps.append(w)

                # Extra settle
                page.wait_for_timeout(800)

                # Pre-actions
                action_warnings = _execute_pre_actions(page, shot, base_url)
                gaps.extend(action_warnings)

                # Hide toasts
                _hide_toasts(page)
                page.wait_for_timeout(300)

                # Capture
                page.screenshot(path=str(out_path), full_page=full_page)
                log.info("  ✅ Saved: %s", out_path.name)

                index.append({
                    "id": sid,
                    "section": shot.get("section", ""),
                    "file": str(out_path.relative_to(REPO_ROOT)),
                    "caption": caption,
                    "url": base_url + url,
                })

            except Exception as exc:
                w = f"Screenshot FAILED for {sid}: {exc}"
                log.error("  ❌ %s", w)
                gaps.append(w)

        ctx.close()
        browser.close()

    return index, gaps

# ─── DOCX assembly ────────────────────────────────────────────────────────────

def _style_exists(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _add_heading(doc: Document, text: str, level: int) -> None:
    style_name = f"Heading {level}"
    if _style_exists(doc, style_name):
        doc.add_heading(text, level=level)
    else:
        p = doc.add_paragraph(text)
        p.style = doc.styles["Normal"]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(16 - (level - 1) * 2)


def _add_body(doc: Document, text: str) -> None:
    style_name = "Body" if _style_exists(doc, "Body") else "Normal"
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph(block, style=style_name)
        p.paragraph_format.space_after = Pt(6)


def _add_screenshot(doc: Document, img_path: Path, caption: str, assets_dir: Path) -> None:
    """Insert a screenshot image with a caption paragraph."""
    if not img_path.exists():
        doc.add_paragraph(f"[Imagen no disponible: {img_path.name}]", style="Normal")
        return

    doc.add_paragraph()  # vertical spacing before image

    # Image paragraph — centered
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    try:
        run.add_picture(str(img_path), width=Inches(5.5))
    except Exception as exc:
        log.warning("Could not add picture %s: %s", img_path.name, exc)
        img_para.add_run(f"[{img_path.name}]")

    # Caption paragraph
    caption_style = "Caption" if _style_exists(doc, "Caption") else "Normal"
    cap_p = doc.add_paragraph(f"Figura: {caption}", style=caption_style)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_style == "Normal":
        for run in cap_p.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # vertical spacing after image


def _add_glossary_table(doc: Document, entries: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Normal Table" if _style_exists(doc, "Normal Table") else "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Término"
    hdr[1].text = "Definición"
    for term, defn in entries:
        row = table.add_row()
        row.cells[0].text = term
        row.cells[1].text = defn


def _add_faq(doc: Document, pairs: list[tuple[str, str]]) -> None:
    for q, a in pairs:
        q_p = doc.add_paragraph(f"P: {q}")
        q_p.style = doc.styles["Normal"]
        for run in q_p.runs:
            run.bold = True
        a_p = doc.add_paragraph(f"R: {a}")
        a_p.style = doc.styles["Normal"]
        a_p.paragraph_format.space_after = Pt(8)


def assemble_docx(
    template_path: Path,
    out_path: Path,
    sections: list[dict],
    screenshot_index: dict[str, dict],
    assets_dir: Path,
) -> list[str]:
    """Build the DOCX manual from the template and content. Returns list of gaps."""
    doc = Document(str(template_path))
    gaps: list[str] = []

    # ── Clear template body content ──────────────────────────────────────────
    # We keep the styles from the template but replace the body content.
    # Remove all paragraphs and tables from the body element.
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sectPr"):
            if tag != "sectPr":  # keep section properties for page layout
                body.remove(child)

    # ── Cover page ───────────────────────────────────────────────────────────
    title_p = doc.add_paragraph("ANCLORA ADVISOR AI")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _style_exists(doc, "Heading 1"):
        title_p.style = doc.styles["Heading 1"]
    for run in title_p.runs:
        run.font.size = Pt(28)
        run.bold = True

    subtitle_p = doc.add_paragraph("Manual de Usuario")
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(4)
    if _style_exists(doc, "Heading 2"):
        subtitle_p.style = doc.styles["Heading 2"]

    date_p = doc.add_paragraph(f"Versión 1.0 — {datetime.now().strftime('%B %Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(36)

    doc.add_page_break()

    # ── Content ──────────────────────────────────────────────────────────────
    for section in sections:
        sid = section.get("id", "")
        heading = section.get("heading", "")
        level = section.get("level", 1)

        _add_heading(doc, heading, level)

        if section.get("content"):
            _add_body(doc, section["content"])

        # Glossary table
        if "glossary" in section:
            _add_glossary_table(doc, section["glossary"])
            doc.add_paragraph()

        # FAQ
        if "faq" in section:
            _add_faq(doc, section["faq"])

        # Subsections
        for sub in section.get("subsections", []):
            _add_heading(doc, sub.get("heading", ""), 2)
            if sub.get("content"):
                _add_body(doc, sub["content"])

            # Screenshot for this subsection
            shot_id = sub.get("screenshot_id")
            if shot_id:
                if shot_id in screenshot_index:
                    meta = screenshot_index[shot_id]
                    img_path = REPO_ROOT / meta["file"]
                    _add_screenshot(doc, img_path, meta["caption"], assets_dir)
                else:
                    w = f"Screenshot '{shot_id}' referenced but not found in index"
                    log.warning(w)
                    gaps.append(w)
                    doc.add_paragraph(f"[Captura no disponible: {shot_id}]", style="Normal")

        # Section-level screenshot (if any)
        shot_id = section.get("screenshot_id")
        if shot_id:
            if shot_id in screenshot_index:
                meta = screenshot_index[shot_id]
                img_path = REPO_ROOT / meta["file"]
                _add_screenshot(doc, img_path, meta["caption"], assets_dir)
            else:
                w = f"Screenshot '{shot_id}' referenced but not found in index"
                gaps.append(w)

        doc.add_paragraph()  # section spacing

    # ── Footer: generation note ───────────────────────────────────────────────
    doc.add_page_break()
    note_p = doc.add_paragraph(
        f"Este manual ha sido generado automáticamente por manual_builder.py "
        f"el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}. "
        "Los elementos marcados como «Pendiente de validación» requieren "
        "revisión manual antes de su publicación."
    )
    note_p.style = doc.styles["Normal"]
    for run in note_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    log.info("DOCX saved: %s", out_path)
    return gaps

# ─── Report ───────────────────────────────────────────────────────────────────

def write_report(
    report_path: Path,
    out_docx: Path,
    screenshot_index: list[dict],
    sections_added: list[str],
    gaps: list[str],
    plan_path: Path,
    template_path: Path,
) -> None:
    glossary_terms = [term for term, _ in MANUAL_SECTIONS[-3].get("glossary", [])]  # type: ignore[index]
    report = {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "generated_files": [str(out_docx.relative_to(REPO_ROOT))],
        "screenshot_index": screenshot_index,
        "sections_added": sections_added,
        "gaps": gaps,
        "sources_used": {
            "template": str(template_path.relative_to(REPO_ROOT)),
            "screenshot_plan": str(plan_path.relative_to(REPO_ROOT)),
            "app_url": APP_URL,
        },
        "glossary": glossary_terms,
        "pending_validation": [
            "Integración VeriFactu directa con AEAT",
            "Importación masiva de facturas",
            "Actualización de contraseña desde perfil",
            "Correo de soporte soporte@anclora.ai",
        ],
    }
    report_path.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    log.info("Report saved: %s", report_path)

# ─── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    _load_env_local()
    args = parse_args()

    plan = load_plan(args.plan)
    all_gaps: list[str] = []
    screenshot_index: list[dict] = []

    # ── Phase 1: Screenshots ──────────────────────────────────────────────────
    if not args.skip_screenshots:
        log.info("=== Phase 1: Taking screenshots ===")
        index, s_gaps = take_screenshots(
            plan=plan,
            auth_state=args.auth_state,
            assets_dir=args.assets_dir,
            headless=args.headless,
        )
        screenshot_index = index
        all_gaps.extend(s_gaps)
        log.info("Screenshots done: %d captured, %d warnings", len(index), len(s_gaps))
    else:
        log.info("Skipping screenshots — using existing files in %s", args.assets_dir)
        # Build index from existing files
        for shot in plan.get("screenshots", []):
            p = args.assets_dir / f"{shot['id']}.png"
            if p.exists():
                screenshot_index.append({
                    "id": shot["id"],
                    "section": shot.get("section", ""),
                    "file": str(p.relative_to(REPO_ROOT)),
                    "caption": shot.get("caption", shot["id"]),
                    "url": plan.get("app_url", APP_URL) + shot.get("url", ""),
                })
            else:
                all_gaps.append(f"Existing screenshot missing: {p}")

    if args.dry_run:
        log.info("Dry-run mode — skipping DOCX assembly.")
        write_report(
            report_path=args.report,
            out_docx=args.out,
            screenshot_index=screenshot_index,
            sections_added=[s["id"] for s in MANUAL_SECTIONS],
            gaps=all_gaps,
            plan_path=args.plan,
            template_path=args.template,
        )
        return

    # ── Phase 2: DOCX assembly ────────────────────────────────────────────────
    log.info("=== Phase 2: Assembling DOCX ===")
    if not args.template.exists():
        log.error("Template not found: %s", args.template)
        sys.exit(1)

    idx_map = {m["id"]: m for m in screenshot_index}
    docx_gaps = assemble_docx(
        template_path=args.template,
        out_path=args.out,
        sections=MANUAL_SECTIONS,
        screenshot_index=idx_map,
        assets_dir=args.assets_dir,
    )
    all_gaps.extend(docx_gaps)

    # ── Phase 3: Report ───────────────────────────────────────────────────────
    write_report(
        report_path=args.report,
        out_docx=args.out,
        screenshot_index=screenshot_index,
        sections_added=[s["id"] for s in MANUAL_SECTIONS],
        gaps=all_gaps,
        plan_path=args.plan,
        template_path=args.template,
    )

    log.info("=== Done ===")
    if all_gaps:
        log.warning("%d gap(s) recorded — see %s", len(all_gaps), args.report)


if __name__ == "__main__":
    main()
